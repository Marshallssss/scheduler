from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path
import smtplib

import pytest
from docx import Document

from scheduler.config import Settings
from scheduler.constants import REPORT_DAILY, REPORT_MONTHLY, REPORT_WEEKLY
from scheduler.repositories import Repository
from scheduler.services.email_service import EmailService
from scheduler.services.progress_service import ProgressService
from scheduler.services.reminder_service import ReminderService
from scheduler.services.report_service import ReportService
from scheduler.utils import is_last_day_of_month
from tests.helpers import FakeEmailService


def test_report_period_windows(session, settings):
    repo = Repository(session)
    svc = ReportService(repo, email_service=FakeEmailService(), report_output_dir=settings.expanded_report_output_dir)

    wk_start, wk_end = svc.period_window(REPORT_WEEKLY, date(2026, 3, 4))
    assert wk_start == date(2026, 3, 2)
    assert wk_end == date(2026, 3, 8)

    mon_start, mon_end = svc.period_window(REPORT_MONTHLY, date(2026, 2, 15))
    assert mon_start == date(2026, 2, 1)
    assert mon_end == date(2026, 2, 28)


def test_last_day_of_month():
    assert is_last_day_of_month(date(2026, 2, 28))
    assert not is_last_day_of_month(date(2026, 2, 27))


def test_email_failure_retries(monkeypatch):
    settings = Settings(
        smtp_host="smtp.example.com",
        smtp_port=587,
        smtp_user="user",
        smtp_pass="pass",
        mail_from="bot@example.com",
    )
    svc = EmailService(settings)

    attempts = {"count": 0}

    def fail_send_once(recipients, subject, body):
        attempts["count"] += 1
        raise RuntimeError("boom")

    monkeypatch.setattr(svc, "_send_once", fail_send_once)
    monkeypatch.setattr("scheduler.services.email_service.time.sleep", lambda _: None)

    ok = svc.send_email(["owner@example.com"], "subject", "body")
    assert not ok
    assert attempts["count"] == 3


def test_email_auth_failure_no_retry(monkeypatch):
    settings = Settings(
        smtp_host="smtp.example.com",
        smtp_port=587,
        smtp_user="user",
        smtp_pass="pass",
        mail_from="bot@example.com",
    )
    svc = EmailService(settings)

    attempts = {"count": 0}

    def fail_auth_once(recipients, subject, body):
        attempts["count"] += 1
        raise smtplib.SMTPAuthenticationError(535, b"5.7.3 Authentication unsuccessful")

    monkeypatch.setattr(svc, "_send_once", fail_auth_once)
    monkeypatch.setattr("scheduler.services.email_service.time.sleep", lambda _: None)

    ok = svc.send_email(["owner@example.com"], "subject", "body")
    assert not ok
    assert attempts["count"] == 1


@pytest.mark.parametrize(
    ("period", "run_date", "report_title"),
    [
        (REPORT_DAILY, date(2026, 3, 2), "项目日报"),
        (REPORT_WEEKLY, date(2026, 3, 6), "项目周报"),
        (REPORT_MONTHLY, date(2026, 3, 31), "项目月报"),
    ],
)
def test_report_includes_goal_charts_for_all_periods(session, settings, period, run_date, report_title):
    base = date(2026, 3, 2)
    repo = Repository(session)

    project = repo.create_project(
        name="Chart Project",
        deadline=base + timedelta(days=15),
        participants=[("Owner", "owner@example.com")],
    )
    phase = repo.add_phase(project.id, name="执行", objective="推进交付")
    owner = repo.list_project_participants(project.id)[0]
    goal = repo.add_goal(
        phase_id=phase.id,
        title="图表目标",
        owner_participant_id=owner.id,
        milestone_date=base + timedelta(days=2),
        deadline=base + timedelta(days=5),
        weight=1,
    )

    progress = ProgressService(repo)
    progress.record_progress(
        goal.id,
        base,
        progress_percent=None,
        requirement_total_count=20,
        requirement_done_count=6,
        progress_state="delayed",
        risk_note="供应商联调延迟",
        updated_by="pm",
    )

    rendered = ReportService(
        repo,
        email_service=FakeEmailService(),
        report_output_dir=settings.expanded_report_output_dir,
    ).render_report(period=period, run_date=run_date)

    assert report_title in rendered.markdown
    assert "目标概览图表" in rendered.markdown
    assert "class=\"report-chart-card\"" in rendered.markdown
    assert "目标进度图" in rendered.markdown
    assert "| 项目 | 阶段 | 目标 | 负责人 | 完成率 | 进度状态 | 权重 | 里程碑 | 截止日期 | 风险项目 |" in rendered.markdown
    assert "<!DOCTYPE html>" in rendered.html
    assert "report-chart-card" in rendered.html
    assert "供应商联调延迟" in rendered.html


def test_report_docx_export_contains_chart_sections(session, settings):
    base = date(2026, 3, 2)
    repo = Repository(session)

    project = repo.create_project(
        name="Chart Project",
        deadline=base + timedelta(days=15),
        participants=[("Owner", "owner@example.com")],
    )
    phase = repo.add_phase(project.id, name="执行", objective="推进交付")
    owner = repo.list_project_participants(project.id)[0]
    goal = repo.add_goal(
        phase_id=phase.id,
        title="图表目标",
        owner_participant_id=owner.id,
        milestone_date=base + timedelta(days=2),
        deadline=base + timedelta(days=5),
        weight=1,
    )

    ProgressService(repo).record_progress(
        goal.id,
        base,
        progress_percent=None,
        requirement_total_count=20,
        requirement_done_count=6,
        progress_state="delayed",
        risk_note="供应商联调延迟",
        updated_by="pm",
    )

    service = ReportService(
        repo,
        email_service=FakeEmailService(),
        report_output_dir=settings.expanded_report_output_dir,
    )
    docx_path = service.export_report_docx(period=REPORT_DAILY, run_date=base)

    assert docx_path.exists()
    assert docx_path.suffix == ".docx"

    document = Document(str(docx_path))
    paragraph_text = "\n".join(item.text for item in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    assert "项目日报" in paragraph_text
    assert "目标概览图表" in paragraph_text
    assert "完成率区间分布" in paragraph_text
    assert "目标状态分布" in paragraph_text
    assert "项目 1: Chart Project" in paragraph_text
    assert "供应商联调延迟" in paragraph_text
    assert "图表目标" in table_text
    assert any(table.rows and len(table.rows[0].cells) == 20 for table in document.tables)


def test_integration_reminder_and_daily_report(session, settings):
    base = date(2026, 3, 2)
    repo = Repository(session)

    project = repo.create_project(
        name="Integration Project",
        deadline=base + timedelta(days=20),
        participants=[("Owner", "owner@example.com")],
    )
    phase = repo.add_phase(project.id, name="P1", objective="Obj")
    owner = repo.list_project_participants(project.id)[0]
    goal = repo.add_goal(
        phase_id=phase.id,
        title="goal",
        owner_participant_id=owner.id,
        milestone_date=base + timedelta(days=1),
        deadline=base + timedelta(days=5),
        weight=2,
    )

    progress = ProgressService(repo)
    progress.record_progress(goal.id, base, 40, updated_by="pm")

    fake_email = FakeEmailService(should_succeed=True)
    reminders = ReminderService(repo, email_service=fake_email, near_days=3)
    report = ReportService(repo, email_service=fake_email, report_output_dir=settings.expanded_report_output_dir)

    reminder_result = reminders.run_milestone_reminders(on_date=base)
    report_result = report.generate_report(period=REPORT_DAILY, run_date=base)

    assert reminder_result.sent == 1
    assert report_result.status == "sent"
    assert Path(report_result.markdown_path).exists()
    assert len(fake_email.sent_messages) >= 2
    assert any(message[3] for message in fake_email.sent_messages if "项目日报" in message[1])
