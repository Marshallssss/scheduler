from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime
from html import escape
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from jinja2 import Environment, FileSystemLoader
from markdown_it import MarkdownIt

from scheduler.constants import (
    EMAIL_STATUS_FAILED,
    EMAIL_STATUS_SENT,
    PROGRESS_STATE_AHEAD,
    PROGRESS_STATE_DELAYED,
    PROGRESS_STATE_NORMAL,
    REPORT_DAILY,
    REPORT_MONTHLY,
    REPORT_PERIODS,
    REPORT_WEEKLY,
)
from scheduler.repositories import GoalSnapshot, Repository
from scheduler.services.email_service import EmailService
from scheduler.utils import month_range, week_range, weighted_progress


@dataclass
class ReportResult:
    markdown_path: Path
    status: str
    report_id: int


@dataclass
class RenderedReport:
    period: str
    run_date: date
    start_date: date
    end_date: date
    subject: str
    markdown: str
    html: str


class ReportService:
    def __init__(self, repo: Repository, email_service: EmailService, report_output_dir: Path) -> None:
        self.repo = repo
        self.email_service = email_service
        self.report_output_dir = report_output_dir

        templates_dir = Path(__file__).resolve().parent.parent / "templates"
        self.env = Environment(
            loader=FileSystemLoader(str(templates_dir)),
            autoescape=False,
            trim_blocks=True,
            lstrip_blocks=True,
        )
        self.markdown_renderer = MarkdownIt("commonmark", {"html": True}).enable("table")

    def period_window(self, period: str, run_date: date) -> tuple[date, date]:
        if period == REPORT_DAILY:
            return run_date, run_date
        if period == REPORT_WEEKLY:
            return week_range(run_date)
        if period == REPORT_MONTHLY:
            return month_range(run_date)
        raise ValueError(f"不支持的报表周期: {period}")

    def generate_report(self, period: str, run_date: date) -> ReportResult:
        return self.dispatch_report(period=period, run_date=run_date)

    def render_report(self, period: str, run_date: date) -> RenderedReport:
        if period not in REPORT_PERIODS:
            raise ValueError(f"不支持的报表周期: {period}")

        start_date, end_date = self.period_window(period, run_date)
        context = self._build_render_context(period=period, start_date=start_date, end_date=end_date, as_of=end_date)
        content = self._render_markdown(period=period, context=context)
        subject = self._subject(period, start_date, end_date)
        return RenderedReport(
            period=period,
            run_date=run_date,
            start_date=start_date,
            end_date=end_date,
            subject=subject,
            markdown=content,
            html=self.render_html_document(content, subject),
        )

    def export_report_docx(self, period: str, run_date: date) -> Path:
        if period not in REPORT_PERIODS:
            raise ValueError(f"不支持的报表周期: {period}")

        start_date, end_date = self.period_window(period, run_date)
        context = self._build_render_context(period=period, start_date=start_date, end_date=end_date, as_of=end_date)
        subject = self._subject(period, start_date, end_date)

        self.report_output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = self.report_output_dir / f"{period}_{start_date}_{end_date}_{timestamp}.docx"
        document = self._build_docx_document(period=period, subject=subject, context=context)
        document.save(str(file_path))
        return file_path

    def export_report_outlook_email(
        self,
        period: str,
        run_date: date,
        recipients: list[str] | None = None,
        markdown: str | None = None,
    ) -> Path:
        rendered = self.render_report(period=period, run_date=run_date)
        content = rendered.markdown if markdown is None else markdown
        html_content = rendered.html if markdown is None else self.render_html_document(content, rendered.subject)

        if recipients is None:
            use_recipients = self.email_service.normalize_recipients(
                [participant.email for participant in self.repo.list_all_participants() if participant.email]
            )
        else:
            use_recipients = self.email_service.normalize_recipients(recipients)

        self.report_output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = self.report_output_dir / f"{period}_{rendered.start_date}_{rendered.end_date}_{timestamp}.eml"
        payload = self.email_service.build_email_bytes(
            recipients=use_recipients,
            subject=rendered.subject,
            body=content,
            html_body=html_content,
            from_address=getattr(getattr(self.email_service, "settings", None), "mail_from", ""),
        )
        file_path.write_bytes(payload)
        return file_path

    def dispatch_report(
        self,
        period: str,
        run_date: date,
        recipients: list[str] | None = None,
        markdown: str | None = None,
    ) -> ReportResult:
        rendered = self.render_report(period=period, run_date=run_date)
        content = rendered.markdown if markdown is None else markdown
        html_content = rendered.html if markdown is None else self.render_html_document(content, rendered.subject)

        self.report_output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = self.report_output_dir / f"{period}_{rendered.start_date}_{rendered.end_date}_{timestamp}.md"
        file_path.write_text(content, encoding="utf-8")

        record = self.repo.create_report_record(
            period=period,
            period_start=rendered.start_date,
            period_end=rendered.end_date,
            markdown_path=str(file_path),
            status="generated",
        )

        use_recipients = recipients
        if use_recipients is None:
            use_recipients = sorted({participant.email for participant in self.repo.list_all_participants()})
        sent = self.email_service.send_email(use_recipients, rendered.subject, content, html_body=html_content)
        status = EMAIL_STATUS_SENT if sent else EMAIL_STATUS_FAILED
        self.repo.mark_report_emailed(record.id, status=status)

        return ReportResult(markdown_path=file_path, status=status, report_id=record.id)

    def render_html_document(self, markdown: str, subject: str) -> str:
        body = self.markdown_renderer.render(markdown)
        title = escape(subject)
        return (
            "<!DOCTYPE html>"
            "<html lang=\"zh-CN\">"
            "<head>"
            "<meta charset=\"utf-8\">"
            "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">"
            f"<title>{title}</title>"
            "<style>"
            "body{margin:0;padding:24px;background:#f5f7fb;color:#132238;font:14px/1.65 -apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;}"
            ".report-doc{max-width:980px;margin:0 auto;background:#fff;border:1px solid #d7deeb;border-radius:20px;box-shadow:0 16px 40px rgba(15,23,42,.08);padding:28px;}"
            "h1,h2,h3{color:#10233d;margin-top:0;}"
            "h1{font-size:28px;margin-bottom:18px;} h2{font-size:22px;margin-top:28px;margin-bottom:14px;} h3{font-size:18px;margin-top:22px;margin-bottom:12px;}"
            "p,li{color:#334155;} ul{padding-left:20px;}"
            "table{width:100%;border-collapse:collapse;margin:14px 0 20px;font-size:13px;}"
            "th,td{border:1px solid #d7deeb;padding:8px 10px;text-align:left;vertical-align:top;}"
            "th{background:#eef4ff;color:#10233d;} tr:nth-child(even) td{background:#fafcff;}"
            ".report-chart-card{margin:12px 0 18px;padding:16px 18px;border:1px solid #d7deeb;border-radius:18px;background:linear-gradient(180deg,#fbfdff 0%,#f5f8ff 100%);}"
            ".report-chart-title{font-size:14px;font-weight:700;color:#10233d;margin-bottom:12px;}"
            ".report-chart-empty{padding:12px;border-radius:14px;background:#fff;color:#64748b;border:1px dashed #cbd5e1;}"
            ".report-badge{display:inline-block;padding:2px 10px;border-radius:999px;font-size:12px;font-weight:700;}"
            ".report-badge.normal{background:#e0f2fe;color:#075985;}"
            ".report-badge.delayed{background:#fee2e2;color:#b91c1c;}"
            ".report-badge.ahead{background:#dcfce7;color:#166534;}"
            ".report-goal-grid{display:grid;gap:10px;}"
            ".report-goal-card{padding:14px;border:1px solid #d7deeb;border-radius:16px;background:#fff;}"
            ".report-goal-meta{font-size:12px;color:#64748b;margin-top:4px;}"
            ".report-goal-risk{margin-top:10px;padding:8px 10px;border-radius:12px;background:#fff1f2;color:#9f1239;font-size:12px;}"
            "</style>"
            "</head>"
            "<body>"
            f"<article class=\"report-doc\">{body}</article>"
            "</body>"
            "</html>"
        )

    def _subject(self, period: str, start_date: date, end_date: date) -> str:
        return f"[{self._report_title(period)}] {start_date.isoformat()} ~ {end_date.isoformat()}"

    def _report_title(self, period: str) -> str:
        labels = {
            REPORT_DAILY: "日报",
            REPORT_WEEKLY: "周报",
            REPORT_MONTHLY: "月报",
        }
        return f"项目{labels[period]}"

    def _build_render_context(self, period: str, start_date: date, end_date: date, as_of: date) -> dict[str, object]:
        grouped = self.repo.grouped_goal_snapshots_by_project(as_of=as_of)
        projects_context = [self._project_context(items) for _, items in sorted(grouped.items(), key=lambda x: x[0])]
        goal_details = [
            goal
            for project in projects_context
            for goal in project["goals"]
        ]

        updates = self.repo.list_progress_updates_between(start_date, end_date)
        rollback_count = sum(1 for update in updates if (update.note or "").strip())
        chart_context = self._goal_chart_context(goal_details)

        return {
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "period": period,
            "start_date": start_date,
            "end_date": end_date,
            "projects": projects_context,
            "goal_details": goal_details,
            "update_count": len(updates),
            "rollback_count": rollback_count,
            **chart_context,
        }

    def _render_markdown(self, period: str, context: dict[str, object]) -> str:
        template = self.env.get_template(f"{period}_report.md.j2")
        return template.render(**context)

    def _project_context(self, snapshots: list[GoalSnapshot]) -> dict:
        project = snapshots[0].project
        weighted_values = [(item.progress, item.goal.weight) for item in snapshots]
        overall_progress = weighted_progress(weighted_values)

        phase_map: dict[int, list[GoalSnapshot]] = defaultdict(list)
        for item in snapshots:
            phase_map[item.phase.id].append(item)

        phases = []
        for phase_id, items in sorted(phase_map.items(), key=lambda pair: pair[0]):
            phase_progress = weighted_progress([(it.progress, it.goal.weight) for it in items])
            phases.append(
                {
                    "id": phase_id,
                    "name": items[0].phase.name,
                    "objective": items[0].phase.objective,
                    "progress": phase_progress,
                    "total_goals": len(items),
                    "completed_goals": sum(1 for it in items if it.progress >= 100),
                }
            )

        goals = [
            {
                "project_id": project.id,
                "project_name": self._markdown_cell(project.name),
                "title": self._markdown_cell(item.goal.title),
                "owner": self._markdown_cell(item.owner.name),
                "owner_email": item.owner.email,
                "progress": item.progress,
                "progress_display": self._progress_display(item.progress, item.progress_state),
                "weight": item.goal.weight,
                "milestone_date": item.goal.milestone_date,
                "deadline": item.goal.deadline,
                "phase_name": self._markdown_cell(item.phase.name),
                "progress_state": item.progress_state,
                "progress_state_label": self._markdown_cell(self._progress_state_label(item.progress_state)),
                "risk_item": self._markdown_cell(self._risk_item(item.progress_state, item.risk_note)),
            }
            for item in snapshots
        ]

        return {
            "id": project.id,
            "name": project.name,
            "deadline": project.deadline,
            "overall_progress": overall_progress,
            "overall_progress_chart_html": self._progress_bar_html(
                label="项目整体完成率",
                progress=overall_progress,
                meta=f"{overall_progress:.2f}% · {sum(1 for item in snapshots if item.progress >= 100)}/{len(snapshots)} 已完成",
                progress_state=PROGRESS_STATE_NORMAL,
            ),
            "total_goals": len(snapshots),
            "completed_goals": sum(1 for item in snapshots if item.progress >= 100),
            "phases": phases,
            "goals": goals,
            "phase_progress_chart_html": self._phase_progress_chart_html(phases),
            "goal_progress_chart_html": self._goal_progress_chart_html(goals),
        }

    def _goal_chart_context(self, goal_details: list[dict]) -> dict[str, object]:
        total = len(goal_details)
        progress_buckets = [
            self._distribution_row("0%", sum(1 for goal in goal_details if float(goal["progress"]) <= 0), total),
            self._distribution_row(
                "1-25%",
                sum(1 for goal in goal_details if 0 < float(goal["progress"]) <= 25),
                total,
            ),
            self._distribution_row(
                "26-50%",
                sum(1 for goal in goal_details if 25 < float(goal["progress"]) <= 50),
                total,
            ),
            self._distribution_row(
                "51-75%",
                sum(1 for goal in goal_details if 50 < float(goal["progress"]) <= 75),
                total,
            ),
            self._distribution_row(
                "76-99%",
                sum(1 for goal in goal_details if 75 < float(goal["progress"]) < 100),
                total,
            ),
            self._distribution_row("100%", sum(1 for goal in goal_details if float(goal["progress"]) >= 100), total),
        ]
        state_rows = [
            self._distribution_row(
                "提前",
                sum(1 for goal in goal_details if goal["progress_state"] == PROGRESS_STATE_AHEAD),
                total,
                progress_state=PROGRESS_STATE_AHEAD,
            ),
            self._distribution_row(
                "正常",
                sum(1 for goal in goal_details if goal["progress_state"] == PROGRESS_STATE_NORMAL),
                total,
            ),
            self._distribution_row(
                "延迟",
                sum(1 for goal in goal_details if goal["progress_state"] == PROGRESS_STATE_DELAYED),
                total,
                progress_state=PROGRESS_STATE_DELAYED,
            ),
        ]
        return {
            "progress_distribution_chart_html": self._distribution_chart_html("完成率区间分布", progress_buckets),
            "state_distribution_chart_html": self._distribution_chart_html("目标状态分布", state_rows),
            "goal_progress_distribution_rows": progress_buckets,
            "goal_state_distribution_rows": state_rows,
        }

    def _distribution_row(
        self,
        label: str,
        count: int,
        total: int,
        progress_state: str = PROGRESS_STATE_NORMAL,
    ) -> dict[str, object]:
        share = (count * 100 / total) if total else 0.0
        return {
            "label": label,
            "count": count,
            "share": share,
            "progress": share,
            "meta": f"{count} 个目标 · {share:.2f}%",
            "progress_state": progress_state,
        }

    def _distribution_chart_html(self, title: str, rows: list[dict[str, object]]) -> str:
        if not rows or not any(int(row["count"]) > 0 for row in rows):
            return (
                f"<div class=\"report-chart-card\">"
                f"<div class=\"report-chart-title\">{escape(title)}</div>"
                "<div class=\"report-chart-empty\">暂无目标数据。</div>"
                "</div>"
            )
        items = "".join(
            self._progress_bar_html(
                label=str(row["label"]),
                progress=float(row["progress"]),
                meta=str(row["meta"]),
                progress_state=str(row["progress_state"]),
            )
            for row in rows
        )
        return (
            f"<div class=\"report-chart-card\">"
            f"<div class=\"report-chart-title\">{escape(title)}</div>"
            f"{items}"
            "</div>"
        )

    def _phase_progress_chart_html(self, phases: list[dict[str, object]]) -> str:
        if not phases:
            return (
                "<div class=\"report-chart-card\">"
                "<div class=\"report-chart-title\">阶段进度图</div>"
                "<div class=\"report-chart-empty\">暂无阶段数据。</div>"
                "</div>"
            )
        rows = "".join(
            self._progress_bar_html(
                label=str(phase["name"]),
                progress=float(phase["progress"]),
                meta=f"{int(phase['completed_goals'])}/{int(phase['total_goals'])} 已完成",
                progress_state=PROGRESS_STATE_NORMAL,
            )
            for phase in phases
        )
        return (
            "<div class=\"report-chart-card\">"
            "<div class=\"report-chart-title\">阶段进度图</div>"
            f"{rows}"
            "</div>"
        )

    def _goal_progress_chart_html(self, goals: list[dict[str, object]]) -> str:
        if not goals:
            return (
                "<div class=\"report-chart-card\">"
                "<div class=\"report-chart-title\">目标进度图</div>"
                "<div class=\"report-chart-empty\">暂无目标数据。</div>"
                "</div>"
            )
        sorted_goals = sorted(
            goals,
            key=lambda goal: (
                0 if goal["progress_state"] == PROGRESS_STATE_DELAYED else 1,
                float(goal["progress"]),
                str(goal["title"]),
            ),
        )
        items = "".join(self._goal_card_html(goal) for goal in sorted_goals)
        return (
            "<div class=\"report-chart-card\">"
            "<div class=\"report-chart-title\">目标进度图</div>"
            f"<div class=\"report-goal-grid\">{items}</div>"
            "</div>"
        )

    def _goal_card_html(self, goal: dict[str, object]) -> str:
        risk_item = str(goal["risk_item"])
        risk_html = ""
        if risk_item != "-":
            risk_html = f"<div class=\"report-goal-risk\">风险提示：{escape(risk_item)}</div>"
        return (
            "<div class=\"report-goal-card\">"
            "<div style=\"display:flex;justify-content:space-between;gap:12px;align-items:flex-start;\">"
            "<div>"
            f"<div style=\"font-weight:700;color:#10233d;\">{escape(str(goal['title']))}</div>"
            f"<div class=\"report-goal-meta\">{escape(str(goal['phase_name']))} · {escape(str(goal['owner']))}</div>"
            f"<div class=\"report-goal-meta\">里程碑 {escape(str(goal['milestone_date']))} · 截止 {escape(str(goal['deadline']))}</div>"
            "</div>"
            f"{self._state_badge_html(str(goal['progress_state']))}"
            "</div>"
            f"{self._progress_bar_html('完成率', float(goal['progress']), str(goal['progress_display']).replace('**', ''), str(goal['progress_state']))}"
            f"{risk_html}"
            "</div>"
        )

    def _progress_bar_html(
        self,
        label: str,
        progress: float,
        meta: str,
        progress_state: str,
    ) -> str:
        normalized = max(0.0, min(100.0, float(progress)))
        fill_color, track_color = self._chart_colors(progress_state)
        return (
            "<div style=\"margin:0 0 12px;\">"
            "<div style=\"display:flex;justify-content:space-between;gap:12px;align-items:center;\">"
            f"<span style=\"font-size:13px;font-weight:600;color:#10233d;\">{escape(label)}</span>"
            f"<span style=\"font-size:12px;color:#64748b;\">{escape(meta)}</span>"
            "</div>"
            f"<div style=\"margin-top:6px;height:12px;background:{track_color};border-radius:999px;overflow:hidden;\">"
            f"<div style=\"height:12px;width:{normalized:.2f}%;background:{fill_color};border-radius:999px;\"></div>"
            "</div>"
            "</div>"
        )

    def _chart_colors(self, progress_state: str) -> tuple[str, str]:
        if progress_state == PROGRESS_STATE_DELAYED:
            return "#ef4444", "#fee2e2"
        if progress_state == PROGRESS_STATE_AHEAD:
            return "#22c55e", "#dcfce7"
        return "#3b82f6", "#dbeafe"

    def _state_badge_html(self, progress_state: str) -> str:
        label = self._progress_state_label(progress_state)
        return f"<span class=\"report-badge {escape(progress_state)}\">{escape(label)}</span>"

    def _progress_state_label(self, progress_state: str) -> str:
        if progress_state == "ahead":
            return "提前"
        if progress_state == "delayed":
            return "delay"
        return "正常"

    def _progress_display(self, progress: float, progress_state: str) -> str:
        value = f"{progress:.2f}%"
        if progress_state == "delayed":
            return f"**{value}**"
        return value

    def _risk_item(self, progress_state: str, risk_note: str | None) -> str:
        clean_note = risk_note.strip() if risk_note else ""
        parts: list[str] = []
        if progress_state == "delayed":
            parts.append("进度delay")
        if clean_note:
            parts.append(clean_note)
        return "；".join(parts) if parts else "-"

    def _markdown_cell(self, value: object) -> str:
        text = str(value).strip() if value is not None else "-"
        if not text:
            return "-"
        return text.replace("|", r"\|").replace("\r\n", "\n").replace("\n", "<br>")

    def _build_docx_document(self, period: str, subject: str, context: dict[str, object]) -> DocxDocument:
        document = Document()
        document.core_properties.title = subject

        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        normal_style = document.styles["Normal"]
        normal_style.font.size = Pt(10.5)

        document.add_heading(self._report_title(period), level=0).alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._add_docx_meta_line(document, "标题", subject)
        self._add_docx_meta_line(document, "生成时间", str(context["generated_at"]))
        if period == REPORT_DAILY:
            self._add_docx_meta_line(document, "周期", str(context["start_date"]))
        else:
            self._add_docx_meta_line(
                document,
                "周期",
                f"{context['start_date']} ~ {context['end_date']}",
            )
        self._add_docx_meta_line(document, "进度更新条数", str(context["update_count"]))
        self._add_docx_meta_line(document, "回退备注条数", str(context["rollback_count"]))

        document.add_heading("目标概览图表", level=1)
        self._add_docx_distribution_section(
            document,
            "完成率区间分布",
            list(context["goal_progress_distribution_rows"]),
        )
        self._add_docx_distribution_section(
            document,
            "目标状态分布",
            list(context["goal_state_distribution_rows"]),
        )

        projects = list(context["projects"])
        if not projects:
            document.add_paragraph("当前无活跃项目。")

        for project in projects:
            document.add_heading(f"项目 {project['id']}: {project['name']}", level=1)
            self._add_docx_meta_line(document, "项目截止日期", str(project["deadline"]))
            self._add_docx_meta_line(document, "总完成率", f"{float(project['overall_progress']):.2f}%")
            self._add_docx_meta_line(
                document,
                "目标完成数",
                f"{int(project['completed_goals'])}/{int(project['total_goals'])}",
            )

            self._add_docx_progress_row(
                document,
                label="项目整体完成率",
                progress=float(project["overall_progress"]),
                meta=(
                    f"{float(project['overall_progress']):.2f}% · "
                    f"{int(project['completed_goals'])}/{int(project['total_goals'])} 已完成"
                ),
                progress_state=PROGRESS_STATE_NORMAL,
            )

            document.add_heading("阶段进度图", level=2)
            phases = list(project["phases"])
            if not phases:
                document.add_paragraph("暂无阶段数据。")
            for phase in phases:
                self._add_docx_progress_row(
                    document,
                    label=self._doc_text(phase["name"]),
                    progress=float(phase["progress"]),
                    meta=f"{int(phase['completed_goals'])}/{int(phase['total_goals'])} 已完成",
                    progress_state=PROGRESS_STATE_NORMAL,
                )

            document.add_heading("目标进度图", level=2)
            goals = sorted(list(project["goals"]), key=self._goal_sort_key)
            if not goals:
                document.add_paragraph("暂无目标数据。")
            for goal in goals:
                self._add_docx_goal_card(document, goal)

        document.add_heading("目标明细（汇总）", level=1)
        self._add_docx_goal_details_table(document, list(context["goal_details"]))
        return document

    def _add_docx_distribution_section(
        self,
        document: DocxDocument,
        title: str,
        rows: list[dict[str, object]],
    ) -> None:
        document.add_heading(title, level=2)
        if not rows or not any(int(row["count"]) > 0 for row in rows):
            document.add_paragraph("暂无目标数据。")
            return
        for row in rows:
            self._add_docx_progress_row(
                document,
                label=self._doc_text(row["label"]),
                progress=float(row["progress"]),
                meta=self._doc_text(row["meta"]),
                progress_state=str(row["progress_state"]),
            )

    def _add_docx_goal_card(self, document: DocxDocument, goal: dict[str, object]) -> None:
        title = document.add_paragraph()
        title.paragraph_format.space_after = Pt(2)
        title_run = title.add_run(self._doc_text(goal["title"]))
        title_run.bold = True
        title.add_run(f"  [{self._doc_text(goal['progress_state_label'])}]")

        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(2)
        meta.add_run(
            f"{self._doc_text(goal['phase_name'])} · {self._doc_text(goal['owner'])}"
        )

        dates = document.add_paragraph()
        dates.paragraph_format.space_after = Pt(2)
        dates.add_run(
            f"里程碑 {self._doc_text(goal['milestone_date'])} · 截止 {self._doc_text(goal['deadline'])}"
        )

        self._add_docx_progress_row(
            document,
            label="完成率",
            progress=float(goal["progress"]),
            meta=self._doc_text(str(goal["progress_display"]).replace("**", "")),
            progress_state=str(goal["progress_state"]),
        )

        risk_item = self._doc_text(goal["risk_item"])
        if risk_item != "-":
            risk = document.add_paragraph()
            risk.paragraph_format.space_after = Pt(8)
            risk_run = risk.add_run(f"风险提示：{risk_item}")
            risk_run.italic = True

    def _add_docx_goal_details_table(
        self,
        document: DocxDocument,
        goal_details: list[dict[str, object]],
    ) -> None:
        headers = ["项目", "阶段", "目标", "负责人", "完成率", "进度状态", "权重", "里程碑", "截止日期", "风险项目"]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            self._set_docx_cell_text(header_cells[index], header, bold=True, fill="DCE6F1")

        rows = goal_details or [
            {
                "project_name": "-",
                "phase_name": "-",
                "title": "-",
                "owner": "-",
                "progress_display": "-",
                "progress_state_label": "-",
                "weight": "-",
                "milestone_date": "-",
                "deadline": "-",
                "risk_item": "-",
            }
        ]
        for goal in rows:
            cells = table.add_row().cells
            values = [
                self._doc_text(goal["project_name"]),
                self._doc_text(goal["phase_name"]),
                self._doc_text(goal["title"]),
                self._doc_text(goal["owner"]),
                self._doc_text(str(goal["progress_display"]).replace("**", "")),
                self._doc_text(goal["progress_state_label"]),
                self._doc_text(f"{goal['weight']:.2f}" if isinstance(goal["weight"], (int, float)) else goal["weight"]),
                self._doc_text(goal["milestone_date"]),
                self._doc_text(goal["deadline"]),
                self._doc_text(goal["risk_item"]),
            ]
            for index, value in enumerate(values):
                self._set_docx_cell_text(cells[index], value)

    def _add_docx_meta_line(self, document: DocxDocument, label: str, value: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(value)

    def _add_docx_progress_row(
        self,
        document: DocxDocument,
        label: str,
        progress: float,
        meta: str,
        progress_state: str,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        title_run = paragraph.add_run(label)
        title_run.bold = True
        paragraph.add_run(f"  {meta}")

        table = document.add_table(rows=1, cols=20)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        fill_color, track_color = self._chart_colors(progress_state)
        filled_cells = int(round(max(0.0, min(100.0, float(progress))) / 5))
        for index, cell in enumerate(table.rows[0].cells):
            cell.width = Inches(0.28)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_docx_cell_text(cell, " ", fill=(fill_color if index < filled_cells else track_color))

        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)

    def _set_docx_cell_text(self, cell, text: str, bold: bool = False, fill: str | None = None) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = bold
        if fill:
            self._set_docx_cell_fill(cell, fill)

    def _set_docx_cell_fill(self, cell, fill: str) -> None:
        color = fill.lstrip("#").upper()
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)

    def _goal_sort_key(self, goal: dict[str, object]) -> tuple[int, float, str]:
        return (
            0 if goal["progress_state"] == PROGRESS_STATE_DELAYED else 1,
            float(goal["progress"]),
            str(goal["title"]),
        )

    def _doc_text(self, value: object) -> str:
        text = str(value).strip() if value is not None else "-"
        if not text:
            return "-"
        return text.replace("<br>", "\n").replace(r"\|", "|")
